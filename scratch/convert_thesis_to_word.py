import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """표 셀의 배경색을 지정하는 헬퍼 함수"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """표 셀의 상하좌우 여백을 조절하여 패딩 추가"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_paragraph_left_border(paragraph, hex_color="002060"):
    """본문 강조 단락(인용문 등)에 왼쪽 테두리를 추가하여 프리미엄 콜아웃 느낌을 제공"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="12" w:color="{hex_color}"/></w:pBdr>')
    pPr.append(pBdr)

def apply_text_formatting(paragraph, text, font_size=Pt(10), font_color=RGBColor(51, 51, 51)):
    """텍스트 내 마크다운 굵게(**text**), 코드블록(`text`) 서식을 변환하여 문단(paragraph)에 추가"""
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            sub_text = part[2:-2]
            new_run = paragraph.add_run(sub_text)
            new_run.font.name = '맑은 고딕'
            new_run.font.size = font_size
            new_run.font.color.rgb = font_color
            new_run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            sub_text = part[1:-1]
            new_run = paragraph.add_run(sub_text)
            new_run.font.name = 'Consolas'
            new_run.font.size = font_size - Pt(0.5)
            new_run.font.color.rgb = RGBColor(192, 0, 0)
        else:
            new_run = paragraph.add_run(part)
            new_run.font.name = '맑은 고딕'
            new_run.font.size = font_size
            new_run.font.color.rgb = font_color

def convert_md_to_docx(md_path, docx_path):
    print(f"📖 논문 마크다운 파일 읽는 중: {md_path}")
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    doc = Document()
    
    # 기본 스타일 여백 정의 (1인치)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # 기본 서체 정의
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10)
    font.color.rgb = RGBColor(51, 51, 51) # 본문: charcoal
    
    in_code_block = False
    in_table = False
    table_headers = []
    table_rows = []
    
    print("✍️ 워드 파일로 파싱 및 변환 작업 중...")
    
    for idx, line in enumerate(lines):
        clean_line = line.strip()
        
        # 1. 아스키 아트 또는 텍스트 기반 다이어그램 코드 블록 처리
        if clean_line.startswith("```"):
            if not in_code_block:
                in_code_block = True
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.left_indent = Inches(0.4)
                add_paragraph_left_border(p, "002060")
                run = p.add_run("📊 [시스템 흐름 아키텍처 다이어그램]\n")
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 32, 96)
                continue
            else:
                in_code_block = False
                continue
                
        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_after = Pt(2)
            add_paragraph_left_border(p, "002060")
            run = p.add_run(line) # 개행 문자 보존을 위해 원본 줄 사용
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(120, 120, 120)
            continue
            
        # 2. 표 (Table) 탐색 및 처리
        if clean_line.startswith("|") and not in_table:
            in_table = True
            table_headers = [cell.strip() for cell in clean_line.split("|")[1:-1]]
            table_rows = []
            continue
        elif in_table and clean_line.startswith("|"):
            if "---" in clean_line:
                continue
            row_data = [cell.strip() for cell in clean_line.split("|")[1:-1]]
            table_rows.append(row_data)
            continue
        elif in_table and not clean_line.startswith("|"):
            in_table = False
            print(f"📊 비교 분석 표 검출 (열: {len(table_headers)}, 행: {len(table_rows)})")
            
            table = doc.add_table(rows=1, cols=len(table_headers))
            table.autofit = True
            
            # 헤더 렌더링
            hdr_cells = table.rows[0].cells
            for col_idx, header_text in enumerate(table_headers):
                hdr_cells[col_idx].text = ""
                set_cell_background(hdr_cells[col_idx], "002060") # 딥 네이비
                set_cell_margins(hdr_cells[col_idx], top=120, bottom=120, left=150, right=150)
                
                p = hdr_cells[col_idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                apply_text_formatting(p, header_text, font_size=Pt(9.5), font_color=RGBColor(255, 255, 255))
            
            # 데이터 행 렌더링
            for r_data in table_rows:
                row_cells = table.add_row().cells
                for col_idx, cell_value in enumerate(r_data):
                    row_cells[col_idx].text = ""
                    set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=150, right=150)
                    
                    if len(table.rows) % 2 == 0:
                        set_cell_background(row_cells[col_idx], "F9F9F9")
                        
                    p = row_cells[col_idx].paragraphs[0]
                    if col_idx == 0:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                    apply_text_formatting(p, cell_value, font_size=Pt(9.5), font_color=RGBColor(51, 51, 51))
            
            doc.add_paragraph().paragraph_format.space_before = Pt(12)
            
        if in_table:
            continue
            
        if not clean_line:
            continue
            
        # 3. 헤더 및 제목 처리
        if clean_line.startswith("# "):
            title_text = clean_line.replace("# 📄 ", "").replace("# ", "").replace("📜 ", "").strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(title_text)
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0, 32, 96)
            run.bold = True
        elif clean_line.startswith("### 부제:"):
            subject_text = clean_line.replace("### ", "").strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(36)
            run = p.add_run(subject_text)
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(89, 89, 89)
            run.bold = True
        elif clean_line.startswith("## "):
            h_text = clean_line.replace("## ", "").strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(h_text)
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0, 32, 96)
            run.bold = True
        elif clean_line.startswith("### "):
            h_text = clean_line.replace("### ", "").strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(h_text)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(46, 117, 181)
            run.bold = True
        elif clean_line.startswith("#### "):
            h_text = clean_line.replace("#### ", "").strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(h_text)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(51, 51, 51)
            run.bold = True
            
        # 4. 콜아웃/인용구 (> )
        elif clean_line.startswith("> "):
            blockquote_text = clean_line.replace("> [!NOTE]", "💡 [논리 정의]").replace("> [!IMPORTANT]", "🔔 [중요 요점]").replace("> ", "").strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.2
            
            apply_text_formatting(p, blockquote_text, font_size=Pt(9.5), font_color=RGBColor(89, 89, 89))
            add_paragraph_left_border(p, "595959")
            
        # 5. 리스트 항목 (* 또는 - 또는 번호 기호)
        elif clean_line.startswith("* ") or clean_line.startswith("- "):
            list_text = clean_line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            apply_text_formatting(p, list_text, font_size=Pt(9.5), font_color=RGBColor(51, 51, 51))
            
        elif re.match(r'^(\d+\.|\w+\)|[①-⑩])\s+', clean_line):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            apply_text_formatting(p, clean_line, font_size=Pt(9.5), font_color=RGBColor(51, 51, 51))
            
        # 6. 일반 본문
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.2
            apply_text_formatting(p, clean_line, font_size=Pt(9.5), font_color=RGBColor(51, 51, 51))

    # 저장
    doc.save(docx_path)
    print(f"🎉 논문 워드 문서 생성 완료: {docx_path}")

if __name__ == "__main__":
    workspace_dir = r"C:\Users\USER\Desktop\Workstation\situation"
    md_file = os.path.join(workspace_dir, "doc", "AI_지능형_운영_시스템_논문.md")
    
    # 1. doc 폴더 내에 지능형_운영_시스템_연구_논문.docx 저장
    docx_file_doc = os.path.join(workspace_dir, "doc", "지능형_운영_시스템_연구_논문.docx")
    convert_md_to_docx(md_file, docx_file_doc)
    
    # 2. 루트 폴더에도 복사하여 찾기 쉽도록 저장
    docx_file_root = os.path.join(workspace_dir, "지능형_운영_시스템_연구_논문.docx")
    convert_md_to_docx(md_file, docx_file_root)
